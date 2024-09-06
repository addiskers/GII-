from pymongo import MongoClient
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# MongoDB connection setup
client = MongoClient('mongodb://localhost:27017/') 
db = client['MnMdata']  
collection = db['structured_headings']  

# Create a new Word document
doc = Document()

# Set default font and size
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Find and process the document
document = collection.find_one()

if document:
    for i in document.get('chapters', []):
        if "BY" in i['chapter'] and "REGION" not in i['chapter'] or "COMPANY" in i["chapter"]:
            
            # Add chapter as a bullet point
            chapter_paragraph = doc.add_paragraph(i['chapter'])
            chapter_paragraph.style = 'List Bullet'
            
            for subsection in i['sub_sections']:
                # Add subsection as an indented bullet point
                subsection_paragraph = doc.add_paragraph(subsection)
                subsection_paragraph.style = 'List Bullet 2'
else:
    doc.add_paragraph("No document found.")

# Save the document
doc.save("output_chapters.docx")
